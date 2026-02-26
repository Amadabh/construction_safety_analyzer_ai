import os
import sys
import json
from pathlib import Path
from collections import Counter
sys.path.insert(0, str(Path(__file__).parent.parent))
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ReportLab — pure Python PDF, no LibreOffice needed
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from schemas import ProcessingResult, AlertLevel, Detection
from config import Config
import re

ALERT_COLORS = {
    AlertLevel.LOW:      ("70AD47", RGBColor(0x70, 0xAD, 0x47)),
    AlertLevel.MEDIUM:   ("FFC000", RGBColor(0xFF, 0xC0, 0x00)),
    AlertLevel.HIGH:     ("FF0000", RGBColor(0xFF, 0x00, 0x00)),
    AlertLevel.CRITICAL: ("7030A0", RGBColor(0x70, 0x30, 0xA0)),
}

# ReportLab color equivalents
ALERT_RL_COLORS = {
    AlertLevel.LOW:      colors.HexColor("#548D2F"),
    AlertLevel.MEDIUM:   colors.HexColor("#D7A91C"),
    AlertLevel.HIGH:     colors.HexColor("#A63030"),
    AlertLevel.CRITICAL: colors.HexColor("#7030A0"),
}

VIOLATION_LABELS = {"NO-Hardhat", "NO-Mask", "NO-Safety Vest"}
MACHINERY_LABELS = {"Excavator", "Wheel Loader", "Machinery", "Dump Truck", "machinery"}

REPORT_SYSTEM_PROMPT = """You are a senior construction site safety officer writing a formal incident report for legal and compliance purposes.

You will be given JSON data containing:
- risk_score and alert_level
- raw_detections: everything the AI vision system detected, with counts and average confidence per label
- violations: the specific PPE/safety violations (subset of raw_detections)
- equipment_on_site: heavy machinery detected
- workers_detected: number of people observed
- applicable_regulations: OSHA/CAL-OSHA regulations retrieved from the compliance database

Write a detailed, professional safety incident report with EXACTLY these sections using markdown headings:

## 1. Executive Summary
2-3 sentences summarizing the overall safety situation, risk level, and most critical findings. Mention specific numbers (e.g. "6 instances of workers without hard hats were detected across 10 frames").

## 2. Site Conditions Overview
Describe what was observed — number of workers, equipment present, general activity level. Reference the specific detection counts from the data.

## 3. Violations Detected
For EACH unique violation type found, write a dedicated paragraph:
- State exactly how many instances were detected and the average confidence score
- Explain in detail why this specific violation is dangerous — what injuries or fatalities can result
- Describe a realistic accident scenario that could occur if this goes uncorrected
- Reference any equipment present that amplifies the risk

## 4. Applicable Regulations
For each regulation:
- State the exact citation
- Explain what the regulation requires workers and employers to do
- Explain specifically how the detected violations breach this regulation
- State what enforcement action OSHA can take

## 5. Risk Assessment Analysis
Explain the risk score in detail:
- Why this specific score was assigned
- How the combination of violations and equipment present elevates risk
- What would need to change to reduce the score

## 6. Recommended Corrective Actions
For each violation type, provide specific actionable steps:
- Immediate (stop-work, PPE distribution, supervisor notification)
- Short-term (toolbox talks, training, accountability measures)
- Long-term (procurement policy, daily PPE checks, monitoring)

## 7. Compliance & Legal Implications
- Potential OSHA citation categories (Serious, Willful, Repeat)
- Penalty ranges under Cal/OSHA
- Liability exposure if a worker is injured under these conditions

Use formal, precise language. Reference actual detection counts and confidence scores throughout."""


def _clean_inline_markdown(text: str) -> str:
    """Remove bold/italic markdown markers from inline text."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # **bold**
    text = re.sub(r'\*(.*?)\*', r'\1', text)        # *italic*
    text = re.sub(r'__(.*?)__', r'\1', text)        # __bold__
    return text

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


class ReportGenerator:
    def __init__(self):
        from model import BedrockModel
        self.model = BedrockModel.get_instance()

    def generate_report(self, result: ProcessingResult, detections: list[Detection] = None) -> str:
        print("Generating report...")
        ra         = result.risk_assessment
        detections = detections or []

        # Build per-label stats from raw detections
        label_stats = {}
        for d in detections:
            if d.label not in label_stats:
                label_stats[d.label] = {"count": 0, "confidence_sum": 0.0}
            label_stats[d.label]["count"]          += 1
            label_stats[d.label]["confidence_sum"] += d.confidence

        raw_detections_summary = [
            {
                "label":              label,
                "count":              stats["count"],
                "avg_confidence":     f"{stats['confidence_sum'] / stats['count']:.1%}",
                "is_violation":       label in VIOLATION_LABELS,
                "is_heavy_equipment": label in MACHINERY_LABELS,
            }
            for label, stats in sorted(label_stats.items())
        ]

        violations_summary = [d for d in raw_detections_summary if d["is_violation"]]
        equipment_summary  = [d for d in raw_detections_summary if d["is_heavy_equipment"]]
        workers_detected   = label_stats.get("Person", {}).get("count", 0)

        context = {
            "video_id":        os.path.basename(result.video_id),
            "analysis_time":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "risk_score":      ra.risk_score,
            "alert_level":     ra.alert_level.value,
            "raw_detections":  raw_detections_summary,
            "violations":      violations_summary,
            "equipment_on_site": equipment_summary,
            "workers_detected":  workers_detected,
            "risk_assessor_violations": [
                {
                    "type":      v.type,
                    "severity":  v.severity.value,
                    "reasoning": getattr(v, "reasoning", ""),
                }
                for v in ra.violations
            ],
            "applicable_regulations": [
                {"citation": r.citation, "full_text": r.text, "source": r.source}
                for r in result.regulations
            ],
        }

        try:
            report_text = self.model.invoke(
                REPORT_SYSTEM_PROMPT,
                json.dumps(context, indent=2)
            )
        except Exception as e:
            print(f"⚠ LLM report generation failed: {e}")
            report_text = self._fallback_report(result, violations_summary)

        safe_name   = os.path.basename(result.video_id).replace(" ", "_")
        base_path   = os.path.join(Config.REPORTS_DIR, safe_name)
        os.makedirs(Config.REPORTS_DIR, exist_ok=True)

        # Build both DOCX and PDF
        docx_path = base_path + "_report.docx"
        pdf_path  = base_path + "_report.pdf"

        self._build_docx(result, report_text, raw_detections_summary, docx_path)
        self._build_pdf(result, report_text, raw_detections_summary, pdf_path)

        print(f"✓ DOCX saved: {docx_path}")
        print(f"✓ PDF saved:  {pdf_path}")
        return docx_path  # return DOCX as primary output (PDF also saved on disk)

    def _fallback_report(self, result: ProcessingResult, violations_summary: list) -> str:
        ra    = result.risk_assessment
        lines = [
            "## 1. Executive Summary",
            f"AI safety analysis detected violations with a risk score of "
            f"{ra.risk_score}/100 (Alert Level: {ra.alert_level.value}).",
            "",
            "## 2. Violations Detected",
        ]
        for v in violations_summary:
            lines.append(f"- {v['label']}: {v['count']} instances "
                         f"(avg confidence: {v['avg_confidence']})")
        lines += [
            "", "## 3. Recommended Actions",
            "Immediately halt work and ensure all workers are equipped with "
            "required PPE before resuming operations."
        ]
        return "\n".join(lines)

    # ── PDF builder (ReportLab) ───────────────────────────────────────────────
    def _build_pdf(self, result: ProcessingResult, report_text: str,
                   raw_detections_summary: list, output_path: str):
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            leftMargin=inch, rightMargin=inch,
            topMargin=inch,  bottomMargin=inch
        )
        ra      = result.risk_assessment
        styles  = getSampleStyleSheet()
        story   = []

        # Custom styles
        title_style = ParagraphStyle(
            "ReportTitle", parent=styles["Title"],
            fontSize=20, spaceAfter=6, textColor=colors.HexColor("#1a1a2e")
        )
        h1_style = ParagraphStyle(
            "H1", parent=styles["Heading1"],
            fontSize=14, spaceBefore=14, spaceAfter=4,
            textColor=colors.HexColor("#2E3A4A"),
            borderPad=4
        )
        body_style = ParagraphStyle(
            "Body", parent=styles["Normal"],
            fontSize=10, spaceAfter=6, leading=14
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=styles["Normal"],
            fontSize=10, spaceAfter=3, leftIndent=20, leading=14,
            bulletIndent=10
        )

        # ── Title ─────────────────────────────────────────────────────────────
        story.append(Paragraph("Construction Site Safety Incident Report", title_style))
        story.append(Paragraph(
            f"<b>Video:</b> {os.path.basename(result.video_id)} &nbsp;&nbsp; "
            f"<b>Generated:</b> {datetime.now().strftime('%B %d, %Y %H:%M:%S')}",
            body_style
        ))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        story.append(Spacer(1, 10))

        # ── Risk score banner ─────────────────────────────────────────────────
        alert_color = ALERT_RL_COLORS.get(ra.alert_level, colors.black)
        banner_data = [[
            Paragraph(f"<b>Risk Score: {ra.risk_score}/100</b>", ParagraphStyle(
                "banner_l", fontSize=14, textColor=alert_color
            )),
            Paragraph(f"<b>Alert Level: {ra.alert_level.value}</b>", ParagraphStyle(
                "banner_r", fontSize=14, textColor=colors.white, alignment=TA_CENTER
            )),
        ]]
        banner_tbl = Table(banner_data, colWidths=[3.5*inch, 3.5*inch])
        banner_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#F5F5F5")),
            ("BACKGROUND", (1, 0), (1, 0), alert_color),
            ("BOX",        (0, 0), (-1, -1), 1, colors.grey),
            ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 18),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 18),
            ("LEFTPADDING",   (0, 0), (-1, -1), 14),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 14),
        ]))
        story.append(banner_tbl)
        story.append(Spacer(1, 16))

        # ── Detection summary table ───────────────────────────────────────────
        if raw_detections_summary:
            story.append(Paragraph("Detection Summary", h1_style))
            tbl_data = [["Label", "Count", "Avg Confidence", "Type"]]
            for d in raw_detections_summary:
                type_label = "VIOLATION" if d["is_violation"] else \
                             "Equipment"  if d["is_heavy_equipment"] else "Person/Object"
                tbl_data.append([
                    d["label"], str(d["count"]), d["avg_confidence"], type_label
                ])

            det_tbl = Table(tbl_data, colWidths=[2.2*inch, 0.8*inch, 1.5*inch, 1.5*inch])
            tbl_style = [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E3A4A")),
                ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",   (0, 0), (-1, -1), 9),
                ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9F9F9")]),
                ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
                ("PADDING",    (0, 0), (-1, -1), 6),
            ]
            # Highlight violation rows
            for i, d in enumerate(raw_detections_summary, start=1):
                if d["is_violation"]:
                    tbl_style.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#FFE0E0")))
                    tbl_style.append(("FONTNAME",   (0, i), (-1, i), "Helvetica-Bold"))

            det_tbl.setStyle(TableStyle(tbl_style))
            story.append(det_tbl)
            story.append(Spacer(1, 16))

        # ── LLM report body ───────────────────────────────────────────────────
        for line in report_text.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            if line.startswith("### "):
                story.append(Paragraph(_clean_inline_markdown(line[4:].strip()), h1_style))
            elif line.startswith("## "):
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
                story.append(Paragraph(_clean_inline_markdown(line[3:].strip()), h1_style))
            elif line.startswith("# "):
                story.append(Paragraph(_clean_inline_markdown(line[2:].strip()), title_style))
            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph(f"• {_clean_inline_markdown(line[2:].strip())}", bullet_style))
            else:
                story.append(Paragraph(_clean_inline_markdown(line), body_style))

        doc.build(story)

    # ── DOCX builder (kept as backup) ────────────────────────────────────────
    def _build_docx(self, result: ProcessingResult, report_text: str,
                    raw_detections_summary: list, output_path: str):
        doc = Document()
        ra  = result.risk_assessment

        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        doc.add_heading("Construction Site Safety Incident Report", 0)
        p = doc.add_paragraph()
        p.add_run("Video File: ").bold = True
        p.add_run(os.path.basename(result.video_id))
        p2 = doc.add_paragraph()
        p2.add_run("Report Generated: ").bold = True
        p2.add_run(datetime.now().strftime("%B %d, %Y %H:%M:%S"))
        doc.add_paragraph()

        hex_color, rgb_color = ALERT_COLORS.get(ra.alert_level, ("000000", RGBColor(0, 0, 0)))
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        cells = tbl.rows[0].cells
        set_cell_bg(cells[0], "F5F5F5")
        r0 = cells[0].paragraphs[0].add_run(f"Risk Score: {ra.risk_score}/100")
        r0.bold = True; r0.font.size = Pt(16); r0.font.color.rgb = rgb_color
        set_cell_bg(cells[1], hex_color)
        r1 = cells[1].paragraphs[0].add_run(f"Alert Level: {ra.alert_level.value}")
        r1.bold = True; r1.font.size = Pt(16)
        r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph()

        if raw_detections_summary:
            doc.add_heading("Detection Summary", level=2)
            dtbl = doc.add_table(rows=1, cols=4)
            dtbl.style = "Table Grid"
            for cell, label in zip(dtbl.rows[0].cells,
                                   ["Label", "Count", "Avg Confidence", "Type"]):
                run = cell.paragraphs[0].add_run(label)
                run.bold = True
                set_cell_bg(cell, "2E3A4A")
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for d in raw_detections_summary:
                row = dtbl.add_row().cells
                row[0].text = d["label"]
                row[1].text = str(d["count"])
                row[2].text = d["avg_confidence"]
                row[3].text = "VIOLATION" if d["is_violation"] else \
                              "Equipment"  if d["is_heavy_equipment"] else "Person/Object"
                if d["is_violation"]:
                    for cell in row:
                        set_cell_bg(cell, "FFE0E0")
            doc.add_paragraph()

        for line in report_text.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith("### "):
                doc.add_heading(_clean_inline_markdown(line[4:].strip()), level=3)
            elif line.startswith("## "):
                doc.add_heading(_clean_inline_markdown(line[3:].strip()), level=2)
            elif line.startswith("# "):
                doc.add_heading(_clean_inline_markdown(line[2:].strip()), level=1)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(_clean_inline_markdown(line[2:].strip()))
            else:
                doc.add_paragraph(_clean_inline_markdown(line))

        doc.save(output_path)